#!/usr/bin/env python3
"""
Assemble the 100 Travel Books manuscript from markdown files into a Word document.
Numbers books sequentially from 01-101 during assembly.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set hyperlink styling
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def parse_inline_formatting(text, paragraph):
    """
    Parse inline markdown formatting (bold, italic, links) and add to paragraph.
    """
    if not text:
        return

    # Pattern to match **bold**, *italic*, and [text](url)
    pattern = r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('[') and '](' in part:
            # Hyperlink
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                url = match.group(2)
                add_hyperlink(paragraph, link_text, url)
        else:
            # Regular text
            paragraph.add_run(part)


def replace_book_number(content, book_number):
    """
    Replace XX or existing numbers with the actual book number in the content.
    Looks for patterns like "# XX —", "# 52 —", "**# [Entry Number] —", etc.
    """
    # Format number as zero-padded (01, 02, ..., 101)
    num_str = str(book_number).zfill(2)

    # Replace various patterns
    patterns = [
        (r'^#\s+XX\s+—', f'# {num_str} —'),  # # XX —
        (r'^#\s+\d+\s+—', f'# {num_str} —'),  # # 52 —
        (r'^#\s+XX\s*:', f'# {num_str}:'),    # # XX:
        (r'^#\s+\d+\s*:', f'# {num_str}:'),   # # 52:
        (r'^\*\*#\s+\[Entry Number\]\s+—', f'# {num_str} —'),  # **# [Entry Number] —
    ]

    for pattern, replacement in patterns:
        content = re.sub(pattern, replacement, content, flags=re.MULTILINE)

    return content


def add_markdown_to_doc(doc, md_content):
    """
    Parse markdown content and add it to the Word document with appropriate formatting.
    """
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines at the start
        if not line:
            i += 1
            continue

        # Headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()

            if level == 1:
                para = doc.add_heading(text, level=1)
                para.runs[0].font.size = Pt(24)
                para.runs[0].font.bold = True
            elif level == 2:
                para = doc.add_heading(text, level=2)
                para.runs[0].font.size = Pt(18)
                para.runs[0].font.bold = True
            elif level == 3:
                para = doc.add_heading(text, level=3)
                para.runs[0].font.size = Pt(14)
                para.runs[0].font.bold = True
            else:
                para = doc.add_heading(text, level=4)
                para.runs[0].font.size = Pt(12)
                para.runs[0].font.bold = True

        # Horizontal rules
        elif line.strip() == '---':
            para = doc.add_paragraph()
            para.add_run('___________________________________________')
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Block quotes
        elif line.startswith('>'):
            quote_text = line.lstrip('>').strip()
            para = doc.add_paragraph()
            para.style = 'Quote'
            parse_inline_formatting(quote_text, para)
            para_format = para.paragraph_format
            para_format.left_indent = Inches(0.5)
            para_format.right_indent = Inches(0.5)
            for run in para.runs:
                run.italic = True

        # Bullet lists
        elif line.startswith('* ') or line.startswith('- '):
            text = line[2:].strip()
            para = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(text, para)

        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            para = doc.add_paragraph(style='List Number')
            parse_inline_formatting(text, para)

        # Regular paragraphs
        else:
            para = doc.add_paragraph()
            parse_inline_formatting(line, para)

        i += 1


def get_book_order():
    """
    Return the ordered list of books by category in chronological order.
    Now includes 101 books total.
    """
    order = {
        '1_first_routes': [
            'the-travels-marco-polo.md',              # 1300
            'the-rihla-ibn-battuta.md',               # 1355
            'the-first-voyage-antonio-pigafetta.md',  # 1525
            'journey-to-western-islands-johnson.md',  # 1775
            'personal-narrative-von-humboldt.md',     # 1814
            'voyage-of-the-beagle-darwin.md',         # 1839
            'incidents-of-travel-in-yucatan-stephens.md', # 1843
            'oregon-trail-parkman.md',                # 1849
            'travels-in-west-africa-mary-kingsley.md', # 1897
            'sailing-alone-joshua-slocum.md',         # 1900
            'blue-nile-moorehead.md'                  # 1962
        ],
        '2_exploration': [
            'south-ernest-shackleton.md',             # 1919
            'worst-journey-cherry-garrard.md',        # 1922
            'libyan-sands-ralph-bagnold.md',          # 1935
            'cruel-way-ella-maillart.md',             # 1947
            'arabian-sands-thesiger.md',              # 1959
            'coming-into-the-country-john-mcphee.md', # 1977
            'snow-leopard-matthiessen.md',            # 1978
            'tracks-robyn-davidson.md',               # 1980
            'arctic-dreams-barry-lopez.md',           # 1986
            'maiden-voyage-tanya-aebi.md',            # 1989
            'terra-incognita-sara-wheeler.md',        # 1996 - EXTRA
            'into-the-wild-krakauer.md',              # 1996
            'cruelest-journey-kira-salak.md'          # 2004
        ],
        '3_slow_routes': [
            'the-narrow-road-basho.md',               # 1689 - EXTRA
            'travels-with-a-donkey-stevenson.md',     # 1879
            'path-to-rome-belloc.md',                 # 1902
            'man-who-walked-colin-fletcher.md',       # 1968
            'as-i-walked-out-lee.md',                 # 1969
            'a-time-of-gifts-fermor.md',              # 1977
            'roads-to-sata-booth.md',                 # 1985
            'chasing-the-monsoon-frater.md',          # 1990
            'roads-to-santiago-nooteboom.md',         # 1992
            'old-ways-macfarlane.md',                 # 2012
            'walking-with-abel-badkhen.md',           # 2015
            'salt-path-raynor-winn.md'                # 2018
        ],
        '4_frontiers': [
            'turkish-embassy-letters-montagu.md',     # 1763
            'ride-to-khiva-fred-burnaby.md',          # 1876
            'unbeaten-tracks-isabella-bird.md',       # 1880
            'travels-in-arabia-deserta-doughty.md',   # 1888
            'my-journey-to-lhasa-davidneel.md',       # 1927
            'valley-of-the-assasins.md',              # 1934
            'road-to-tartary-fleming.md',             # 1935 (News from Tartary)
            'road-to-oxiana-byron.md',                # 1937
            'in-an-antique-land-ghosh.md',            # 1992
            'from-holy-mountain-dalrymple.md',        # 1997
            'unexpected-light-jason-elliott.md',      # 1999
            'travels-with-a-tangerine-mackintosh-smith.md', # 2001
            'great-hedge-of-india-moxham.md',         # 2001
            'places-in-between-rory-stewart.md'       # 2006
        ],
        '5_humbled': [
            'short-walk-down-hindu-kush-newby.md',    # 1958
            'from-heaven-lake-vikram-seth.md',        # 1983
            'great-american-bus-adler.md',            # 1986
            'video-night-pico-iyer.md',               # 1988
            'walk-in-the-woods-bryson.md',            # 1998
            'dark-continent-khumalo.md',              # 2007
            'turn-right-at-machu-adams.md'            # 2011
        ],
        '6_institutional': [
            'way-of-the-world-bouvier.md',            # 1963
            'full-tilt-devla-murphy.md',              # 1965
            'great-railway-bazaar-theroux.md',        # 1975
            'jupiters-travels-ted-simon.md',          # 1979 - EXTRA
            'blue-highways-heat-moon.md',             # 1982
            'into-the-heart-of-borneo-ohanlon.md',    # 1984
            'motoring-with-md-eric-hansen.md',        # 1991
            'river-center-world-winchester.md',       # 1996
            'fortune-teller-told-terzani.md',         # 1997
            'passage-to-juneau-raban.md',             # 1999
            'shadow-of-silk-road.md',                 # 2006
            'first-overland-tim-slessor.md'           # 2006
        ],
        '7_inhabited_landscapes': [
            'house-in-bali-mcphee.md',                # 1946
            'pattern-of-islands-grimble.md',          # 1952
            'windward-road-archie-carr.md',           # 1956
            'venice-jan-morris.md',                   # 1960
            'cities-and-stones-aldiss.md',            # 1966
            'iron-and-silk-salzman.md',               # 1986
            'year-in-provence-peter-mayle.md',        # 1989
            'river-town-hessler.md',                  # 2001
            'maximum-city-mehta.md',                  # 2004
            'indonesia-etc-pisani.md'                 # 2014
        ],
        '8_political': [
            'black-lamb-rebecca-west.md',             # 1941
            'naples-44-norman-lewis.md',              # 1978
            'africa-calliope-hoagland.md',            # 1979
            'small-place-jamaica-kincaid.md',         # 1988
            'great-game-peter-hopkirk.md',            # 1990
            'my-traitors-heart-malan.md',             # 1990
            'shadow-of-sun-kapuscinski.md',           # 2001
            'when-crocodile-godwin.md',               # 2006
            'blood-river-tim-butcher.md',             # 2008
            'wild-coast-john-gimlette.md',            # 2011
            'border-kapka-kassova.md'                 # 2017
        ],
        '9_self_reckoning': [
            'west-with-the-night-markham.md',         # 1942
            'travels-with-charley-steinbeck.md',      # 1962
            'area-of-darkness-naipaul.md',            # 1964
            'in-patagonia-chatwin.md',                # 1977
            'travels-with-myself-gellhorn.md',        # 1978
            'there-and-then-thubron.md',              # 1992 (James Salter)
            'motorcyle-diaries-guevara.md',           # 1993
            'out-of-sheer-rage-dyer.md',              # 1997
            'catfish-and-mandala-pham.md',            # 1999
            'one-day-i-will-write-wainaina.md',       # 2011
            'looking-for-transwonderland-sarowiwa.md' # 2012
        ]
    }
    return order


def main():
    """
    Main function to assemble the manuscript.
    """
    base_path = Path('/home/ram/projects/100TravelBooks')

    # Create a new Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)

    print("Starting manuscript assembly...")

    # 1. Add Preface
    print("Adding preface...")
    preface_path = base_path / 'introduction' / 'preface_unified.md'
    if preface_path.exists():
        with open(preface_path, 'r', encoding='utf-8') as f:
            add_markdown_to_doc(doc, f.read())
        doc.add_page_break()
    else:
        print(f"Warning: {preface_path} not found")

    # 2. Add Organizing Principle Essay
    print("Adding organizing principle essay...")
    org_principle_path = base_path / 'introduction' / 'organizing_principle_essay.md'
    if org_principle_path.exists():
        with open(org_principle_path, 'r', encoding='utf-8') as f:
            add_markdown_to_doc(doc, f.read())
        doc.add_page_break()
    else:
        print(f"Warning: {org_principle_path} not found")

    # 3. Process each category
    book_order = get_book_order()
    categories = [
        '1_first_routes',
        '2_exploration',
        '3_slow_routes',
        '4_frontiers',
        '5_humbled',
        '6_institutional',
        '7_inhabited_landscapes',
        '8_political',
        '9_self_reckoning'
    ]

    book_number = 1  # Start numbering from 1

    for category in categories:
        print(f"\nProcessing category: {category}")
        category_path = base_path / 'book-writeups' / category

        # Add category introduction
        category_intro_path = category_path / 'category_intro.md'
        if category_intro_path.exists():
            print(f"  Adding category intro...")
            with open(category_intro_path, 'r', encoding='utf-8') as f:
                add_markdown_to_doc(doc, f.read())
        else:
            print(f"  Warning: {category_intro_path} not found")

        # Add each book in chronological order
        if category in book_order:
            for book_file in book_order[category]:
                book_path = category_path / book_file
                if book_path.exists():
                    print(f"  Adding book #{book_number:02d}: {book_file}")
                    with open(book_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        # Replace XX or existing number with current book number
                        content = replace_book_number(content, book_number)
                        add_markdown_to_doc(doc, content)
                    doc.add_paragraph()  # Add spacing between books
                    book_number += 1
                else:
                    print(f"  Warning: {book_path} not found")

        # Add page break after each category
        doc.add_page_break()

    # Save the document
    output_path = base_path / '100-travel-books-manuscript.docx'
    doc.save(output_path)
    print(f"\n✓ Manuscript saved to: {output_path}")
    print(f"✓ Total books numbered: {book_number - 1}")
    print("Done!")


if __name__ == '__main__':
    main()
